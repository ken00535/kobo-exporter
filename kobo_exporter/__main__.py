import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING
from kobo_exporter.model.model import Handler


doc = docx.Document()
handler = Handler()

bookmarks = handler.read_bookmarks()
for bookmark in bookmarks:
    text = doc.add_paragraph(handler.format_bookamrk(bookmark))
    text.style.font.size = Pt(10)
    text.style.font.name = u'新細明體'
    text.style._element.rPr.rFonts.set(qn('w:eastAsia'), u'新細明體')
    text.style.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
    text.paragraph_format.space_before = Pt(0)
    text.paragraph_format.space_after = Pt(0)
    text.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    annotation = handler.format_annotation(bookmark)
    if annotation != '':
        run = text.add_run(handler.format_annotation(bookmark))
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    text = doc.add_paragraph('')
    text.paragraph_format.space_before = Pt(0)
    text.paragraph_format.space_after = Pt(0)
    text.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


doc.save('note.docx')
